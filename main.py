from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Create document
doc = Document()

# Title section
doc.add_paragraph("Numon Adizov", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Address: Karaulbazar, Bukhara, Uzbekistan", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Phone: +998 99 772 66 46", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER
doc.add_paragraph("Email: nomonadizov4@gmail.com", style='Normal').alignment = WD_ALIGN_PARAGRAPH.CENTER

# Add spacing
doc.add_paragraph("\n")

# Add title
title = doc.add_paragraph("Personal Statement for GTE Requirement")
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.runs[0].bold = True

# Content text
content = """
Dear Visa Officer,

My name is Numon Adizov, a 21-year-old citizen of Uzbekistan. I am writing to formally apply for a student visa (subclass 500) to pursue the Master of Information Technology (083790M) with a specialisation in Information Systems at Macquarie University, Sydney, Australia. Through this statement, I aim to demonstrate my genuine intention to study in Australia, outline the reasons for choosing Macquarie University, and explain how this program aligns with my academic background, professional ambitions, and Uzbekistan’s long-term technological goals.

From an early age, I was captivated by the digital world and how technology transforms daily life. Growing up in Karaulbazar, I often spent hours exploring computers, video games, and gadgets. I vividly remember playing GTA as a child, not just for entertainment but for curiosity—questioning how such games were created, drawing rough algorithms on paper, and even noticing bugs within them. This early fascination gradually evolved into a passion for understanding how systems work and how software can shape modern societies. Encouraged by my parents, who recognised my enthusiasm, I began studying programming fundamentals and developing a long-term goal of becoming an IT professional who contributes to digital innovation in Uzbekistan.

After completing my secondary education, I was admitted to one of the top academic lyceums in Uzbekistan—Karakul Academic Lyceum—where I specialised in exact sciences such as mathematics, physics, and computer science. My strong academic performance there helped me gain admission offers from eight universities worldwide, including institutions in the United States, the United Kingdom, South Korea, and Uzbekistan. However, I chose to study at Ajou University in Tashkent (AUT)—a Korean university branch—because of my lifelong admiration for Korean culture, history, and language, which share structural similarities with Uzbek. I also appreciated that my major, Electrical and Computer Engineering, was strongly connected to IT.

During my time at Ajou University (2021–2025), I developed a solid academic foundation in both engineering and computer science. I actively engaged in classes related to software engineering and programming languages such as Python, Kotlin, Java, and C, and I was particularly drawn to Python for its simplicity and its central role in Artificial Intelligence, Data Analytics, and Machine Learning. In my AI class, I built a model capable of identifying traffic signs with high precision—an experience that strengthened my analytical and practical skills. My parents also encouraged me to study Python more deeply at Najot Ta’lim, one of Uzbekistan’s leading IT learning centres. Later, I supplemented my knowledge with online courses on Udemy, CodeWithMosh, HackerRank, and LeetCode, where I earned certificates and improved my problem-solving abilities.

Beyond academics, I engaged in several projects that enriched my leadership and teamwork experience. As the Technical Manager at RoboSmart LLC, a robotics company in Tashkent, I led a small team of university students in creating innovative projects such as a smart-city model, a drone system, and smart-home prototypes. This internship not only honed my managerial and technical abilities but also strengthened my capacity for effective communication and project organisation. Additionally, I volunteered at national-level events such as the FIFA Futsal World Cup 2024 and Maker Faire Uzbekistan, gaining experience in event management and teamwork in multicultural environments.

At Ajou University, I was also the captain of the university intellectual team “Ajou Zakovat Club”, a member of the Ajou Tech Club, and part of Google Developer Group (GDG) on Campus as a technical manager. These experiences deepened my technical expertise and leadership confidence. I graduated with an excellent academic record, scoring 93% out of 100, reflecting consistent effort and commitment to excellence.

Why Australia

I have chosen Australia because of its world-class education system, strong research culture, and welcoming environment for international students. The Australian education system competes with British and American standards while offering more inclusive support and practical experience opportunities. Moreover, Australia is known for its high level of safety, clean environment, and pleasant climate, all of which make it a healthy and inspiring place to study. As an English-speaking country, it also allows me to integrate easily without facing language barriers.

Why Macquarie University

I chose Macquarie University because it perfectly aligns with my academic interests and career goals. Ranked 138th globally by QS World University Rankings, Macquarie is internationally recognised for excellence in information technology and business-oriented programs. What impressed me most was the university’s focus on combining theory with hands-on learning through industry partnerships and internships with leading IT companies in Sydney’s technology corridor.

The Master of Information Technology (specialisation in Information Systems) is designed for students who wish to gain deep technical knowledge alongside business-driven IT management skills. This combination will allow me to design and manage systems that improve business efficiency—skills essential for my future startup projects. Additionally, Macquarie’s modern facilities, sports centres, and multicultural campus will provide me with a balanced academic and personal experience. The university’s Early Acceptance Scholarship has further motivated me, as it reflects recognition of my academic merit and dedication.

Why Not Uzbekistan

Although Uzbekistan’s higher education system is improving, there are still challenges: few universities rank within the world’s top 1000, research facilities remain limited, and many qualified professors lack international teaching exposure. Most major universities are located in Tashkent, leading to overcrowding and limited research infrastructure. While I am proud of Uzbekistan’s progress, I strongly believe that studying in Australia will expose me to cutting-edge IT knowledge and global practices that are not yet available domestically.

Future Goals and Contribution to Uzbekistan

Uzbekistan is rapidly advancing toward digital transformation through initiatives like Digital Uzbekistan 2030, which aims to establish a modern, innovation-driven economy. The country is investing heavily in ICT infrastructure and e-governance, creating a growing demand for globally trained professionals. I want to be part of this transformation by returning home with advanced technical and managerial skills.

My short-term goal after graduation is to work in major IT companies such as EPAM Systems Uzbekistan or in leading digital solution firms contributing to software development, ERP system management, and AI-based solutions. My long-term goal is to launch two startups:
1. Developing an ERP system for large distributors, initially tested in our family’s company to improve efficiency and digital operations.
2. Creating a green-tech platform that promotes environmental awareness by rewarding citizens who take eco-friendly actions—supporting Uzbekistan’s vision of sustainable growth.

Through these initiatives, I aim to promote innovation, sustainability, and digital literacy in Uzbekistan.

Financial Capacity and Family Ties

My studies and living expenses will be fully sponsored by my brother’s private enterprise, “Qumri Adiz Mujizasi”, which has operated successfully for nearly 20 years and generates an annual turnover of over 4 billion UZS. The company is well-known in our region for its trade activities and stable growth. In addition, with my family’s support, we have accumulated USD 57,300 in our bank account dedicated to my education in Australia. The total estimated expenses include AUD 94,500 for tuition and approximately AUD 24,505 per year for living costs, which are fully covered.

Our family is deeply rooted in Uzbekistan, both personally and professionally. My brother and I also plan to expand the company’s activities upon my return, integrating technology into trade operations. These strong family and professional connections ensure my commitment to returning home after completing my studies.

International Exposure and Personal Development

I have previously travelled to Vietnam and the United Arab Emirates, experiences that broadened my global perspective. In Vietnam, I enjoyed the serenity of nature and observed the people’s discipline in maintaining health and fitness. In the UAE, I was inspired by the blend of traditional culture and modern innovation, and I was impressed by the kindness and helpfulness of the local people. These experiences strengthened my adaptability, independence, and appreciation for global diversity—qualities that will help me integrate into Australia’s multicultural society.

Outside my studies, I plan to engage in various extracurricular activities at Macquarie University, including joining the football club and the Sydney Running Club, continuing my passion for fitness. I also plan to run a YouTube and Telegram blog sharing my student life in Australia and exploring academic environments such as libraries and cafés. Additionally, I plan to organise a chess tournament for Uzbek students studying in Australia to foster community and cultural connection.

Conclusion

In conclusion, my decision to study in Australia is driven by a clear academic purpose and a well-defined professional vision that aligns with Uzbekistan’s national development goals. I am confident that Macquarie University’s world-class education and Australia’s inclusive learning environment will provide the ideal foundation for my growth as a professional and a person. Upon completion of my studies, I am fully committed to returning to Uzbekistan to contribute to the country’s ICT development and help drive its digital transformation.

Thank you for considering my application.

Sincerely,
Numon Adizov
"""

doc.add_paragraph(content)

# Save the document
output_path = "Numon_Adizov_GTE_Final.docx"
doc.save(output_path)

output_path
